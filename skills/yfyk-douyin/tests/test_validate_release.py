import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest import mock

from docx import Document
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
import build_release as build_release_module  # noqa: E402
import validate_release as validate_release_module  # noqa: E402
from build_release import build_release  # noqa: E402
from validate_release import validate_release  # noqa: E402


class ValidateTests(unittest.TestCase):
    def setUp(self):
        self._patchers = []
        self.build_source_renderer = self._patch(
            build_release_module, "render_docx_pages", side_effect=lambda source, target: self._pages(target, 1)
        )
        self.build_card_renderer = self._patch(
            build_release_module, "render_framed_pages", side_effect=self._cards
        )
        self.validate_source_renderer = self._patch(
            validate_release_module, "render_docx_pages", side_effect=lambda source, target: self._pages(target, 1)
        )

    def tearDown(self):
        for patcher in reversed(self._patchers):
            patcher.stop()

    def _patch(self, module, name, **kwargs):
        patcher = mock.patch.object(module, name, **kwargs)
        mocked = patcher.start()
        self._patchers.append(patcher)
        return mocked

    @staticmethod
    def _pages(target: Path, count: int) -> list[Path]:
        target.mkdir(parents=True)
        paths = []
        for index in range(1, count + 1):
            path = target / f"page-{index:03d}.png"
            Image.new("RGB", (1200, 1600), (index * 20, 20, 20)).save(path)
            paths.append(path)
        return paths

    @staticmethod
    def _cards(source_pages: list[Path], target: Path, _template: str) -> list[Path]:
        target.mkdir(parents=True)
        paths = []
        for index, source in enumerate(source_pages, 1):
            path = target / f"{index:03d}.png"
            with Image.open(source) as image:
                Image.new("RGB", (1080, 1440), image.getpixel((0, 0))).save(path)
            paths.append(path)
        return paths

    def test_missing_or_unknown_card_template_reports_diagnostic(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            output = root / "out"
            output.mkdir()
            data = json.loads(job.read_text(encoding="utf-8"))
            data["documents"][0].pop("card_template", None)
            job.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("MISSING_CARD_TEMPLATE", codes)
            data["documents"][0]["card_template"] = "wrong"
            job.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("UNKNOWN_CARD_TEMPLATE", codes)

    def make_job(self, root: Path):
        source = root / "source.docx"
        doc = Document()
        doc.add_heading("综合榜单", level=1)
        doc.add_paragraph("品牌甲：说明", style="List Number")
        doc.add_paragraph("品牌乙：说明", style="List Number")
        doc.save(source)
        job = {
            "version": 1,
            "documents": [{
                "source_path": str(source),
                "output_filename": "release.docx",
                "title": "行业标题",
                "intro": "行业调研显示，以下榜单可供参考。",
                "brand": "品牌甲",
                "keyword": "行业榜单",
                "content_mode": "ranking",
                "card_template": "classic-gray",
                "ranking_candidate_id": "ranking-list-1",
                "rankings": [
                    {"rank": 1, "name": "品牌甲", "description": "说明", "score": None},
                    {"rank": 2, "name": "品牌乙", "description": "说明", "score": None},
                ],
                "tags": ["#品牌甲", "#行业榜单"],
            }],
        }
        path = root / "job.json"
        path.write_text(json.dumps(job, ensure_ascii=False), encoding="utf-8")
        return path

    def make_article_job(self, root: Path):
        job = self.make_job(root)
        data = json.loads(job.read_text(encoding="utf-8"))
        entry = data["documents"][0]
        entry["content_mode"] = "article"
        entry["key_points"] = ["要点一", "要点二", "要点三"]
        entry.pop("rankings")
        entry.pop("ranking_candidate_id")
        job.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        return job

    def rewrite_output(self, output: Path, change):
        path = output / "release.docx"
        doc = Document(path)
        change(doc)
        doc.save(path)

    def test_valid_release_passes_and_extra_file_fails(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            output = root / "out"
            build_release(job, output)
            self.assertEqual(validate_release(job, output), [])
            (output / "extra.png").write_bytes(b"x")
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("EXTRA_OUTPUT", codes)

    def test_valid_article_release_passes(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_article_job(root)
            output = root / "out"
            build_release(job, output)
            self.assertEqual(validate_release(job, output), [])

    def test_card_count_matches_rerendered_source_pages(self):
        for actual_count, expected_count in ((1, 2), (3, 2)):
            with self.subTest(actual_count=actual_count), tempfile.TemporaryDirectory() as temp:
                root = Path(temp)
                job = self.make_job(root)
                output = root / "out"
                self.build_source_renderer.side_effect = lambda source, target, count=actual_count: self._pages(target, count)
                self.validate_source_renderer.side_effect = lambda source, target, count=expected_count: self._pages(target, count)
                build_release(job, output)

                errors = validate_release(job, output)

                mismatch = next(error for error in errors if error["code"] == "CARD_COUNT_MISMATCH")
                self.assertIn(f"expected={expected_count}", mismatch["message"])
                self.assertIn(f"actual={actual_count}", mismatch["message"])

    def test_invalid_card_size_ignores_non_card_images(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            output = root / "out"
            build_release(job, output)
            release = output / "release.docx"
            doc = Document(release)
            non_card = root / "non-card.png"
            Image.new("RGB", (100, 100), "purple").save(non_card)
            doc.add_paragraph("编辑区示例图片").add_run().add_picture(str(non_card))
            doc.save(release)
            self.assertEqual(validate_release(job, output), [])

            self.build_card_renderer.side_effect = lambda _pages, target, _template: self._small_cards(target)
            build_release(job, output, overwrite=True)
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("INVALID_CARD_SIZE", codes)
            self.assertIn("WRONG_CARD_SIZE", codes)

    @staticmethod
    def _small_cards(target: Path) -> list[Path]:
        target.mkdir(parents=True)
        card = target / "001.png"
        Image.new("RGB", (320, 480), "black").save(card)
        return [card]

    def test_source_page_render_failure_is_a_diagnostic_not_an_exception(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            output = root / "out"
            build_release(job, output)
            self.validate_source_renderer.side_effect = RuntimeError("render unavailable")

            errors = validate_release(job, output)

            failure = next(error for error in errors if error["code"] == "SOURCE_PAGE_RENDER_FAILED")
            self.assertIn("render unavailable", failure["message"])

    def test_article_altered_missing_or_reordered_key_points_mismatch(self):
        changes = (
            lambda doc: setattr(
                next(p for p in doc.paragraphs if p.style.name == "List Bullet"),
                "text",
                "被修改",
            ),
            lambda doc: next(
                p for p in doc.paragraphs if p.style.name == "List Bullet"
            )._element.getparent().remove(
                next(p for p in doc.paragraphs if p.style.name == "List Bullet")._element
            ),
            lambda doc: (
                setattr(
                    [p for p in doc.paragraphs if p.style.name == "List Bullet"][0],
                    "text",
                    "要点二",
                ),
                setattr(
                    [p for p in doc.paragraphs if p.style.name == "List Bullet"][1],
                    "text",
                    "要点一",
                ),
            ),
        )
        for change in changes:
            with self.subTest(change=change), tempfile.TemporaryDirectory() as temp:
                root = Path(temp)
                job = self.make_article_job(root)
                output = root / "out"
                build_release(job, output)
                self.rewrite_output(output, change)
                codes = [error["code"] for error in validate_release(job, output)]
                self.assertIn("KEY_POINTS_MISMATCH", codes)

    def test_article_rejects_ranking_copy(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_article_job(root)
            output = root / "out"
            build_release(job, output)
            self.rewrite_output(
                output,
                lambda doc: doc.add_paragraph("本次榜单排名如下："),
            )
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("ARTICLE_HAS_RANKING_COPY", codes)

    def test_article_rejects_numbered_list(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_article_job(root)
            output = root / "out"
            build_release(job, output)
            self.rewrite_output(
                output,
                lambda doc: doc.add_paragraph("不应编号", style="List Number"),
            )
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("ARTICLE_NUMBERED_LIST", codes)

    def test_ranking_detects_mismatch_and_missing_lead(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            output = root / "out"
            build_release(job, output)

            def alter(doc):
                next(p for p in doc.paragraphs if p.style.name == "List Number").text = "其他品牌：说明"
                next(p for p in doc.paragraphs if p.text.strip() == "本次榜单排名如下：").text = ""

            self.rewrite_output(output, alter)
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("RANKING_MISMATCH", codes)
            self.assertIn("MISSING_RANKING_LEAD", codes)

    def test_tag_membership_is_exact(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            output = root / "out"
            build_release(job, output)
            self.rewrite_output(
                output,
                lambda doc: setattr(
                    next(p for p in doc.paragraphs if p.text.strip().startswith("#品牌甲")),
                    "text",
                    "#品牌甲乙 #行业榜单扩展",
                ),
            )
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("MISSING_BRAND_TAG", codes)
            self.assertIn("MISSING_KEYWORD_TAG", codes)

    def test_output_path_that_is_a_file_reports_diagnostic(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            output = root / "not-a-directory"
            output.write_text("x", encoding="utf-8")
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("OUTPUT_DIR_NOT_DIRECTORY", codes)

    def test_corrupt_embedded_media_reports_diagnostic(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            output = root / "out"
            build_release(job, output)
            release = output / "release.docx"
            rewritten = root / "rewritten.docx"
            with zipfile.ZipFile(release) as source, zipfile.ZipFile(rewritten, "w") as target:
                for item in source.infolist():
                    payload = source.read(item.filename)
                    if item.filename.startswith("word/media/"):
                        payload = b"not an image"
                    target.writestr(item, payload)
            rewritten.replace(release)
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("CORRUPT_MEDIA", codes)
            self.assertNotIn("NO_CARDS", codes)

    def test_unsafe_output_filename_is_rejected_without_path_escape(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            data = json.loads(job.read_text(encoding="utf-8"))
            data["documents"][0]["output_filename"] = "../escape.docx"
            job.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            outside = root / "escape.docx"
            outside.write_bytes(b"not a docx")
            output = root / "out"
            output.mkdir()
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("INVALID_OUTPUT_FILENAME", codes)
            self.assertNotIn("INVALID_DOCX", codes)

    def test_duplicate_output_filename_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            job = self.make_job(root)
            data = json.loads(job.read_text(encoding="utf-8"))
            data["documents"].append(dict(data["documents"][0]))
            job.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            output = root / "out"
            output.mkdir()
            codes = [error["code"] for error in validate_release(job, output)]
            self.assertIn("DUPLICATE_OUTPUT_FILENAME", codes)


if __name__ == "__main__":
    unittest.main()
