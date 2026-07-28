import sys
import tempfile
import unittest
from pathlib import Path

from PIL import Image
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
from inspect_sources import inspect_document  # noqa: E402


class InspectorTests(unittest.TestCase):
    def test_plain_article_uses_article_mode(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "article.docx"
            doc = Document()
            doc.add_heading("装修经验分享", level=1)
            doc.add_paragraph("装修前应先明确预算，再根据家庭需求规划空间。")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertEqual(result["content_mode"], "article")

    def test_ranked_source_uses_ranking_mode(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "ranking.docx"
            doc = Document()
            doc.add_heading("装修品牌综合榜单", level=1)
            doc.add_paragraph("品牌甲：稳定", style="List Number")
            doc.add_paragraph("品牌乙：节能", style="List Number")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertEqual(result["content_mode"], "ranking")

    def test_article_mode_does_not_warn_about_missing_ranking(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "article.docx"
            doc = Document()
            doc.add_paragraph("这是一篇普通的装修知识文章。")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertNotIn("NO_RANKING", [warning["code"] for warning in result["warnings"]])

    def test_article_infers_repeated_brand_candidate(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "article.docx"
            doc = Document()
            doc.add_heading("装修选择指南", level=1)
            doc.add_paragraph("尚海整装提供从设计到施工的一体化服务。")
            doc.add_paragraph("不少业主选择尚海整装，是因为看重规范的施工流程。")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertEqual(result["brand_candidates"][0]["value"], "尚海整装")
            self.assertIn("confidence", result["brand_candidates"][0])
            self.assertIn("reason", result["brand_candidates"][0])

    def test_article_does_not_infer_generic_home_improvement_phrase(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "generic.docx"
            doc = Document()
            doc.add_paragraph("我们家装时先确定了整体预算。")
            doc.add_paragraph("后来我们家装也优先考虑了收纳需求。")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertEqual(result["brand_candidates"], [])

    def test_article_does_not_infer_brand_mentioned_in_only_one_paragraph(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "single-mention.docx"
            doc = Document()
            doc.add_paragraph("尚海整装提供从设计到施工的一体化服务，尚海整装也重视施工规范。")
            doc.add_paragraph("业主还应结合预算和实际需求进行选择。")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertEqual(result["brand_candidates"], [])

    def test_article_brand_candidates_follow_first_text_position(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "two-brands.docx"
            doc = Document()
            doc.add_paragraph("先了解聚通装饰，再比较尚海整装。")
            doc.add_paragraph("聚通装饰和尚海整装都有各自的服务特点。")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertEqual(
                [candidate["value"] for candidate in result["brand_candidates"]],
                ["聚通装饰", "尚海整装"],
            )

    def test_article_does_not_infer_repeated_generic_brand_shaped_phrases(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "generic-phrases.docx"
            doc = Document()
            doc.add_paragraph("进行家装前要明确预算，整体家装也要考虑收纳。现代家居强调舒适。")
            doc.add_paragraph("进行家装需要规划工期，整体家装需要统一风格。现代家居也重视实用。")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertEqual(result["brand_candidates"], [])

    def test_numbered_list_ranking_keeps_order(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "sample.docx"
            doc = Document()
            doc.add_heading("2026家电综合榜单", level=1)
            doc.add_paragraph("品牌甲：稳定", style="List Number")
            doc.add_paragraph("品牌乙：节能", style="List Number")
            doc.save(source)
            result = inspect_document(source, root / "work")
            names = [item["name"] for item in result["ranking_candidates"][0]["items"]]
            self.assertEqual(names, ["品牌甲", "品牌乙"])

    def test_table_ranking_keeps_score_and_description(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "table.docx"
            doc = Document()
            table = doc.add_table(rows=3, cols=4)
            for cell, value in zip(table.rows[0].cells, ["排名", "品牌", "评分", "说明"]):
                cell.text = value
            for cell, value in zip(table.rows[1].cells, ["1", "品牌甲", "97.5", "稳定"]):
                cell.text = value
            for cell, value in zip(table.rows[2].cells, ["2", "品牌乙", "92", "节能"]):
                cell.text = value
            doc.save(source)
            result = inspect_document(source, root / "work")
            item = result["ranking_candidates"][0]["items"][0]
            self.assertEqual((item["name"], item["score"], item["description"]), ("品牌甲", "97.5", "稳定"))

    def test_pipe_table_paragraphs_are_detected(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "pipe.docx"
            doc = Document()
            doc.add_paragraph("| 排名 | 品牌 | 评分 | 说明 |")
            doc.add_paragraph("|---|---|---:|---|")
            doc.add_paragraph("| 1 | 品牌甲 | 97.5 | 稳定 |")
            doc.add_paragraph("| 2 | 品牌乙 | 92 | 节能 |")
            doc.save(source)
            result = inspect_document(source, root / "work")
            items = result["ranking_candidates"][0]["items"]
            self.assertEqual([(i["name"], i["score"]) for i in items], [("品牌甲", "97.5"), ("品牌乙", "92")])

    def test_multiple_rankings_are_reported_as_ambiguous(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "multi.docx"
            doc = Document()
            for heading, names in (("综合榜单", ["甲", "乙"]), ("推荐榜单", ["丙", "丁"])):
                doc.add_heading(heading, level=1)
                for name in names:
                    doc.add_paragraph(name, style="List Number")
            doc.save(source)
            result = inspect_document(source, root / "work")
            self.assertIn("AMBIGUOUS_RANKING", [warning["code"] for warning in result["warnings"]])

    def test_blocks_keep_paragraph_image_order(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            image_path = root / "source.png"
            Image.new("RGB", (300, 200), "blue").save(image_path)
            source = root / "image.docx"
            doc = Document()
            doc.add_paragraph("图片之前")
            doc.add_picture(str(image_path))
            doc.add_paragraph("图片之后")
            doc.save(source)
            result = inspect_document(source, root / "work")
            kinds = [block["type"] for block in result["blocks"]]
            self.assertEqual(kinds, ["paragraph", "image", "paragraph"])


if __name__ == "__main__":
    unittest.main()
