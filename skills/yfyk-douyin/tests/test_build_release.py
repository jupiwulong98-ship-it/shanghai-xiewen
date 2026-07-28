import json
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
import build_release as build_release_module  # noqa: E402
from build_release import build_docx, build_release, render_cards, validate_job, wrap_text  # noqa: E402


class BuildTests(unittest.TestCase):
    def test_card_template_is_required_and_must_be_registered(self):
        entry = {
            "source_path": "/tmp/source.docx",
            "output_filename": "release.docx",
            "title": "标题",
            "intro": "导语",
            "brand": "品牌甲",
            "keyword": "行业观察",
            "tags": ["#品牌甲", "#行业观察"],
            "content_mode": "article",
            "key_points": ["一", "二", "三"],
        }
        with self.assertRaisesRegex(ValueError, "card_template"):
            validate_job({"version": 1, "documents": [entry]})
        entry["card_template"] = "not-registered"
        with self.assertRaisesRegex(ValueError, "card_template"):
            validate_job({"version": 1, "documents": [entry]})

    def test_editable_copy_uses_a_libreoffice_visible_cjk_font(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            card = root / "card.png"
            Image.new("RGB", (1080, 1440), "white").save(card)
            target = root / "release.docx"
            entry = {
                "title": "中文标题",
                "intro": "中文正文",
                "content_mode": "article",
                "key_points": ["中文列表一", "中文列表二", "中文列表三"],
                "tags": ["#中文"],
            }

            build_docx(entry, [card], target)

            result = Document(target)
            expected_font = "Arial Unicode MS"
            self.assertEqual(build_release_module.WORD_FONT, expected_font)

            def assert_cjk_font(rfonts):
                for font_kind in ("eastAsia", "ascii", "hAnsi"):
                    self.assertEqual(rfonts.get(qn(f"w:{font_kind}")), expected_font)

            assert_cjk_font(result.styles["Normal"]._element.rPr.rFonts)
            normal_paragraph = next(p for p in result.paragraphs if p.text == "中文正文")
            list_paragraph = next(p for p in result.paragraphs if p.text == "中文列表一")
            self.assertEqual(list_paragraph.style.name, "List Bullet")
            assert_cjk_font(normal_paragraph.runs[0]._element.rPr.rFonts)
            assert_cjk_font(list_paragraph.runs[0]._element.rPr.rFonts)

    def test_cards_are_1080_by_1440_and_long_text_uses_multiple_cards(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            text = "完整正文" * 3000
            cards = render_cards(
                [{"type": "paragraph", "style": "Normal", "text": text}],
                root / "cards",
                "classic-gray",
            )
            self.assertGreater(len(cards), 1)
            sizes = []
            for card in cards:
                with Image.open(card) as image:
                    sizes.append(image.size)
            self.assertTrue(all(size == (1080, 1440) for size in sizes))

    def test_wrap_text_preserves_every_character(self):
        image = Image.new("RGB", (400, 200))
        draw = ImageDraw.Draw(image)
        text = "第一行 ABC 123，第二行也不能丢。"
        lines = wrap_text(draw, text, font=ImageFont.load_default(), max_width=50)
        self.assertEqual("".join(lines), text)

    def test_release_docx_has_editable_copy_and_embedded_cards(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            doc = Document()
            doc.add_heading("综合榜单", level=1)
            doc.add_paragraph("品牌甲：说明", style="List Number")
            doc.add_paragraph("品牌乙：说明", style="List Number")
            doc.save(source)
            job = {
                "version": 1,
                "documents": [{
                    "source_path": str(source),
                    "output_filename": "release.docx",
                    "title": "行业榜单标题",
                    "intro": "行业调研显示，以下榜单可供参考。",
                    "brand": "品牌甲",
                    "keyword": "行业榜单",
                    "content_mode": "ranking",
                    "card_template": "classic-gray",
                    "ranking_candidate_id": "ranking-list-1",
                    "rankings": [
                        {"rank": 1, "name": "品牌甲", "description": "说明", "score": None},
                        {"rank": 2, "name": "品牌乙", "description": "说明", "score": None},
                    ],
                    "tags": ["#品牌甲", "#行业榜单"],
                }],
            }
            job_path = root / "job.json"
            job_path.write_text(json.dumps(job, ensure_ascii=False), encoding="utf-8")
            outputs = build_release(job_path, root / "out")
            result = Document(outputs[0])
            texts = [paragraph.text for paragraph in result.paragraphs]
            page_break_index = next(
                index for index, paragraph in enumerate(result.paragraphs)
                if any(
                    br.get(qn("w:type")) == "page"
                    for br in paragraph._element.iter(qn("w:br"))
                )
            )
            self.assertEqual(
                texts[1:page_break_index],
                [
                    "行业榜单标题",
                    "正文",
                    "行业调研显示，以下榜单可供参考。",
                    "本次榜单排名如下：",
                    "品牌甲：说明",
                    "品牌乙：说明",
                    "Tag",
                    "#品牌甲 #行业榜单",
                ],
            )
            ranking_paragraphs = [
                paragraph for paragraph in result.paragraphs
                if paragraph.text in {"品牌甲：说明", "品牌乙：说明"}
            ]
            self.assertEqual(
                [paragraph.style.name for paragraph in ranking_paragraphs],
                ["List Number", "List Number"],
            )
            self.assertNotIn("榜单", texts)
            self.assertEqual(len(result.inline_shapes), 1)

    def test_article_release_has_bulleted_key_points_and_embedded_cards(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            doc = Document()
            doc.add_heading("行业观察", level=1)
            doc.add_paragraph("完整来源内容")
            doc.save(source)
            points = ["趋势正在形成", "用户更看重体验", "服务能力成为关键"]
            job = {
                "version": 1,
                "documents": [{
                    "source_path": str(source),
                    "output_filename": "article.docx",
                    "title": "行业观察标题",
                    "intro": "最新报告显示，行业正在发生变化。",
                    "brand": "品牌甲",
                    "keyword": "行业观察",
                    "tags": ["#品牌甲", "#行业观察"],
                    "content_mode": "article",
                    "card_template": "classic-gray",
                    "key_points": points,
                }],
            }
            job_path = root / "job.json"
            job_path.write_text(json.dumps(job, ensure_ascii=False), encoding="utf-8")

            outputs = build_release(job_path, root / "out")

            result = Document(outputs[0])
            texts = [paragraph.text for paragraph in result.paragraphs]
            self.assertIn("核心要点", texts)
            self.assertNotIn("本次榜单排名如下：", texts)
            for point in points:
                paragraph = next(p for p in result.paragraphs if p.text == point)
                self.assertEqual(paragraph.style.name, "List Bullet")
            self.assertEqual(len(result.inline_shapes), 1)

    def test_three_cards_use_one_section_and_explicit_page_breaks(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            cards = []
            for index in range(1, 4):
                card = root / f"card-{index}.png"
                Image.new("RGB", (1080, 1440), "white").save(card)
                cards.append(card)
            target = root / "release.docx"
            entry = {
                "title": "中文标题",
                "intro": "中文正文",
                "content_mode": "article",
                "key_points": ["要点一", "要点二", "要点三", "要点四"],
                "tags": ["#中文", "#三卡测试"],
            }

            build_docx(entry, cards, target)

            result = Document(target)
            body = result._element.body
            self.assertEqual(len(list(body.iter(qn("w:sectPr")))), 1)
            self.assertEqual(len(result.inline_shapes), 3)
            self.assertEqual(
                [paragraph.text for paragraph in result.paragraphs if paragraph.text.startswith("图 ")],
                ["图 1/3", "图 2/3", "图 3/3"],
            )

            children = list(body)
            card_sequences = []
            for index, child in enumerate(children):
                page_breaks = list(child.iter(qn("w:br")))
                if any(br.get(qn("w:type")) == "page" for br in page_breaks):
                    card_sequences.append(children[index:index + 3])
            self.assertEqual(len(card_sequences), 3)
            for index, (page_break, picture, caption) in enumerate(card_sequences, 1):
                self.assertEqual(page_break.tag, qn("w:p"))
                self.assertEqual(page_break.text, "")
                self.assertEqual(len(list(page_break.iter(qn("w:br")))), 1)
                self.assertEqual(len(list(picture.iter(qn("w:drawing")))), 1)
                self.assertEqual(
                    "".join(node.text or "" for node in caption.iter(qn("w:t"))),
                    f"图 {index}/3",
                )

    def test_job_validation_rejects_invalid_content_modes_and_fields(self):
        common = {
            "source_path": "/tmp/source.docx",
            "output_filename": "release.docx",
            "title": "标题",
            "intro": "导语",
            "brand": "品牌甲",
            "keyword": "行业观察",
            "tags": ["#品牌甲", "#行业观察"],
            "card_template": "classic-gray",
        }
        invalid_entries = [
            ({**common}, "content_mode"),
            ({**common, "content_mode": "unknown"}, "content_mode"),
            ({**common, "content_mode": "ranking", "rankings": []}, "rankings"),
            ({
                **common,
                "content_mode": "ranking",
                "rankings": [{"rank": 1, "name": "品牌甲", "description": "", "score": None}],
                "key_points": [],
            }, "key_points"),
            ({**common, "content_mode": "article", "key_points": ["一", "二"]}, "3"),
            ({**common, "content_mode": "article", "key_points": ["一", "二", "三", "四", "五", "六"]}, "5"),
            ({**common, "content_mode": "article", "key_points": ["一", "", "三"]}, "non-empty"),
            ({**common, "content_mode": "article", "key_points": ["一", 2, "三"]}, "strings"),
            ({**common, "content_mode": "article", "key_points": ["一", "二", "三"], "rankings": []}, "rankings"),
            ({
                **common,
                "content_mode": "article",
                "key_points": ["一", "二", "三"],
                "ranking_candidate_id": "ranking-1",
            }, "ranking_candidate_id"),
        ]
        for entry, message in invalid_entries:
            with self.subTest(entry=entry, message=message):
                with self.assertRaisesRegex(ValueError, message):
                    validate_job({"version": 1, "documents": [entry]})

    def test_ranking_mode_requires_explicit_mode_and_keeps_numbered_list(self):
        entry = {
            "source_path": "/tmp/source.docx",
            "output_filename": "release.docx",
            "title": "标题",
            "intro": "导语",
            "brand": "品牌甲",
            "keyword": "行业榜单",
            "tags": ["#品牌甲", "#行业榜单"],
            "card_template": "classic-gray",
            "content_mode": "ranking",
            "rankings": [{"rank": 1, "name": "品牌甲", "description": "说明", "score": None}],
        }
        validate_job({"version": 1, "documents": [entry]})
        without_mode = dict(entry)
        without_mode.pop("content_mode")
        with self.assertRaisesRegex(ValueError, "content_mode"):
            validate_job({"version": 1, "documents": [without_mode]})

    def test_ranking_and_tags_require_strict_shapes(self):
        common = {
            "source_path": "/tmp/source.docx",
            "output_filename": "release.docx",
            "title": "标题",
            "intro": "导语",
            "brand": "品牌甲",
            "keyword": "行业榜单",
            "content_mode": "ranking",
            "tags": ["#品牌甲", "#行业榜单"],
            "card_template": "classic-gray",
        }
        valid_item = {"rank": 1, "name": "品牌甲", "description": "说明", "score": None}
        invalid = [
            ({**common, "rankings": ["not-a-dict"]}, "dict"),
            ({**common, "rankings": [{**valid_item, "rank": 0}]}, "positive int"),
            ({**common, "rankings": [{**valid_item, "rank": True}]}, "positive int"),
            ({**common, "rankings": [{**valid_item, "name": " "}]}, "name"),
            ({**common, "rankings": [{**valid_item, "description": None}]}, "description"),
            ({**common, "rankings": [{**valid_item, "score": []}]}, "score"),
            ({**common, "rankings": [valid_item], "tags": "#品牌甲 #行业榜单"}, "tags"),
            ({**common, "rankings": [valid_item], "tags": ["#品牌甲", " "]}, "tags"),
        ]
        for entry, message in invalid:
            with self.subTest(message=message):
                with self.assertRaisesRegex(ValueError, message):
                    validate_job({"version": 1, "documents": [entry]})

    def test_required_tags_use_exact_membership(self):
        entry = self._ranking_entry(Path("/tmp/source.docx"), "release.docx")
        entry["tags"] = ["#品牌甲其他", "#行业榜单"]
        with self.assertRaisesRegex(ValueError, "Tag"):
            validate_job({"version": 1, "documents": [entry]})

    def test_duplicate_output_names_are_rejected_before_writing(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = self._make_source(root / "source.docx")
            entries = [
                self._ranking_entry(source, "same.docx"),
                self._ranking_entry(source, "same.docx"),
            ]
            job = self._write_job(root, entries)
            with self.assertRaisesRegex(ValueError, "duplicate"):
                build_release(job, root / "out", overwrite=True)
            self.assertFalse((root / "out" / "same.docx").exists())

    def test_later_collision_leaves_no_earlier_new_output(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = self._make_source(root / "source.docx")
            output = root / "out"
            output.mkdir()
            (output / "existing.docx").write_bytes(b"existing")
            job = self._write_job(root, [
                self._ranking_entry(source, "new.docx"),
                self._ranking_entry(source, "existing.docx"),
            ])
            with self.assertRaises(FileExistsError):
                build_release(job, output)
            self.assertFalse((output / "new.docx").exists())
            self.assertEqual((output / "existing.docx").read_bytes(), b"existing")

    def test_bad_later_source_leaves_no_output(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = self._make_source(root / "source.docx")
            job = self._write_job(root, [
                self._ranking_entry(source, "new.docx"),
                self._ranking_entry(root / "missing.docx", "missing.docx"),
            ])
            with self.assertRaisesRegex(ValueError, "source_path"):
                build_release(job, root / "out")
            self.assertFalse((root / "out" / "new.docx").exists())

    def test_publish_failure_rolls_back_new_and_overwritten_outputs(self):
        for overwrite in (False, True):
            with self.subTest(overwrite=overwrite), tempfile.TemporaryDirectory() as temp:
                root = Path(temp)
                source = self._make_source(root / "source.docx")
                output = root / "out"
                output.mkdir()
                originals = {}
                if overwrite:
                    originals = {"one.docx": b"original-one", "two.docx": b"original-two"}
                    for filename, content in originals.items():
                        (output / filename).write_bytes(content)
                job = self._write_job(root, [
                    self._ranking_entry(source, "one.docx"),
                    self._ranking_entry(source, "two.docx"),
                ])
                real_publish_replace = build_release_module._publish_replace
                publish_count = 0

                def fail_second_staged_publish(source_path, target_path):
                    nonlocal publish_count
                    if source_path.parent.name == "staged":
                        publish_count += 1
                        if publish_count == 2:
                            raise OSError("injected publish failure")
                    real_publish_replace(source_path, target_path)

                with mock.patch.object(
                    build_release_module,
                    "_publish_replace",
                    side_effect=fail_second_staged_publish,
                ):
                    with self.assertRaisesRegex(OSError, "injected"):
                        build_release(job, output, overwrite=overwrite)

                if overwrite:
                    for filename, content in originals.items():
                        self.assertEqual((output / filename).read_bytes(), content)
                else:
                    self.assertFalse((output / "one.docx").exists())
                    self.assertFalse((output / "two.docx").exists())

    def test_overwrite_rejects_non_file_target_before_any_publication(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = self._make_source(root / "source.docx")
            output = root / "out"
            protected_dir = output / "release.docx"
            protected_dir.mkdir(parents=True)
            protected_content = protected_dir / "keep.txt"
            protected_content.write_bytes(b"keep-me")
            job = self._write_job(root, [
                self._ranking_entry(source, "new.docx"),
                self._ranking_entry(source, "release.docx"),
            ])

            with self.assertRaisesRegex(ValueError, "regular file"):
                build_release(job, output, overwrite=True)

            self.assertTrue(protected_dir.is_dir())
            self.assertEqual(protected_content.read_bytes(), b"keep-me")
            self.assertFalse((output / "new.docx").exists())

    @staticmethod
    def _make_source(path: Path) -> Path:
        doc = Document()
        doc.add_paragraph("来源正文")
        doc.save(path)
        return path

    @staticmethod
    def _ranking_entry(source: Path, filename: str):
        return {
            "source_path": str(source),
            "output_filename": filename,
            "title": "标题",
            "intro": "导语",
            "brand": "品牌甲",
            "keyword": "行业榜单",
            "content_mode": "ranking",
            "card_template": "classic-gray",
            "rankings": [{"rank": 1, "name": "品牌甲", "description": "说明", "score": None}],
            "tags": ["#品牌甲", "#行业榜单"],
        }

    @staticmethod
    def _write_job(root: Path, entries: list[dict]) -> Path:
        path = root / "job.json"
        path.write_text(json.dumps({"version": 1, "documents": entries}, ensure_ascii=False), encoding="utf-8")
        return path


if __name__ == "__main__":
    unittest.main()
