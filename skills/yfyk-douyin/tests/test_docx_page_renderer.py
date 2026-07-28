import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch
import importlib.util
import shutil
import stat

from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from docx_page_renderer import (  # noqa: E402
    PageRenderError,
    _find_soffice,
    render_docx_pages,
)


class _FakePixmap:
    def __init__(self, color: str, image_format: str = "PNG", fail: bool = False):
        self.color = color
        self.image_format = image_format
        self.fail = fail

    def save(self, path: str):
        if self.fail:
            raise OSError("simulated page write failure")
        Image.new("RGB", (32, 48), self.color).save(path, format=self.image_format)


class _FakePage:
    def __init__(self, color: str, image_format: str = "PNG", fail: bool = False):
        self.color = color
        self.image_format = image_format
        self.fail = fail

    def get_pixmap(self, *, matrix, alpha):
        self.matrix = matrix
        self.alpha = alpha
        return _FakePixmap(self.color, self.image_format, self.fail)


class _FakeDocument:
    def __init__(self, colors: list[str], image_format: str = "PNG", fail_on_page: int | None = None):
        self.pages = [
            _FakePage(color, image_format, fail=index == fail_on_page)
            for index, color in enumerate(colors)
        ]
        self.page_count = len(self.pages)

    def load_page(self, index: int):
        return self.pages[index]

    def close(self):
        pass


class _FakeFitz:
    class Matrix:
        def __init__(self, x, y):
            self.x = x
            self.y = y

    def __init__(self, colors: list[str], image_format: str = "PNG", fail_on_page: int | None = None):
        self.colors = colors
        self.image_format = image_format
        self.fail_on_page = fail_on_page

    def open(self, path):
        return _FakeDocument(self.colors, self.image_format, self.fail_on_page)


class DocxPageRendererTests(unittest.TestCase):
    def _source(self, root: Path) -> Path:
        source = root / "source.docx"
        source.write_bytes(b"not-a-real-docx")
        return source

    def _successful_conversion(self, args, **kwargs):
        output_dir = Path(args[args.index("--outdir") + 1])
        source = Path(args[-1])
        (output_dir / f"{source.stem}.pdf").write_bytes(b"%PDF-fake")
        return subprocess.CompletedProcess(args, 0, "", "")

    def test_renders_two_pages_in_order_and_as_decodable_pngs(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            target = root / "pages"
            with (
                patch("docx_page_renderer._load_fitz", return_value=_FakeFitz(["red", "blue"])),
                patch("docx_page_renderer.subprocess.run", side_effect=self._successful_conversion),
                patch("docx_page_renderer._find_soffice", return_value="/fake/soffice"),
            ):
                pages = render_docx_pages(self._source(root), target, dpi=160)

            self.assertEqual(pages, [target / "page-001.png", target / "page-002.png"])
            self.assertEqual(sorted(path.name for path in target.glob("page-*.png")), ["page-001.png", "page-002.png"])
            for page in pages:
                with Image.open(page) as image:
                    image.verify()
                self.assertEqual(stat.S_IMODE(page.stat().st_mode), 0o644)

    def test_missing_source_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp:
            missing = Path(temp) / "missing.docx"
            with self.assertRaisesRegex(PageRenderError, "does not exist"):
                render_docx_pages(missing, Path(temp) / "pages")

    def test_conversion_failure_and_zero_page_pdf_are_rejected(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = self._source(root)
            with (
                patch("docx_page_renderer._find_soffice", return_value="/fake/soffice"),
                patch(
                    "docx_page_renderer.subprocess.run",
                    return_value=subprocess.CompletedProcess([], 1, "", "converter failed"),
                ),
            ):
                with self.assertRaisesRegex(PageRenderError, "conversion failed"):
                    render_docx_pages(source, root / "conversion-failure")

            with (
                patch("docx_page_renderer._find_soffice", return_value="/fake/soffice"),
                patch("docx_page_renderer.subprocess.run", side_effect=self._successful_conversion),
                patch("docx_page_renderer._load_fitz", return_value=_FakeFitz([])),
            ):
                with self.assertRaisesRegex(PageRenderError, "no pages"):
                    render_docx_pages(source, root / "zero-pages")

    def test_existing_target_directory_is_rejected_before_rendering(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            target = root / "pages"
            target.mkdir()
            with self.assertRaisesRegex(PageRenderError, "target directory already exists"):
                render_docx_pages(self._source(root), target)

    def test_jpeg_disguised_as_png_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            with (
                patch("docx_page_renderer._load_fitz", return_value=_FakeFitz(["red"], "JPEG")),
                patch("docx_page_renderer.subprocess.run", side_effect=self._successful_conversion),
                patch("docx_page_renderer._find_soffice", return_value="/fake/soffice"),
            ):
                with self.assertRaisesRegex(PageRenderError, "not a decodable PNG"):
                    render_docx_pages(self._source(root), root / "pages")

    def test_later_page_failure_leaves_no_target_directory(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            target = root / "pages"
            with (
                patch("docx_page_renderer._load_fitz", return_value=_FakeFitz(["red", "blue"], fail_on_page=1)),
                patch("docx_page_renderer.subprocess.run", side_effect=self._successful_conversion),
                patch("docx_page_renderer._find_soffice", return_value="/fake/soffice"),
            ):
                with self.assertRaisesRegex(PageRenderError, "failed to rasterize"):
                    render_docx_pages(self._source(root), target)
            self.assertFalse(target.exists())

    def _pdftoppm_run(self, mode: str):
        def fake_run(args, **kwargs):
            if "--convert-to" in args:
                return self._successful_conversion(args, **kwargs)
            if mode == "nonzero":
                return subprocess.CompletedProcess(args, 1, "", "rasterizer failed")
            if mode == "timeout":
                raise subprocess.TimeoutExpired(args, 120)
            prefix = Path(args[-1])
            if mode == "gap":
                page_numbers = [1, 3]
            elif mode == "empty":
                page_numbers = []
            else:
                page_numbers = [1, 2]
            for number in page_numbers:
                Image.new("RGB", (20, 30), "navy").save(f"{prefix}-{number}.png", format="PNG")
            if mode == "stray":
                (prefix.parent / "unexpected.txt").write_text("unexpected")
            return subprocess.CompletedProcess(args, 0, "", "")

        return fake_run

    def test_pdftoppm_fallback_renames_multiple_pages_and_sets_permissions(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            target = root / "pages"
            with (
                patch("docx_page_renderer._load_fitz", side_effect=PageRenderError("fitz missing")),
                patch("docx_page_renderer._find_pdftoppm", return_value="/fake/pdftoppm"),
                patch("docx_page_renderer._find_soffice", return_value="/fake/soffice"),
                patch("docx_page_renderer.subprocess.run", side_effect=self._pdftoppm_run("success")),
            ):
                pages = render_docx_pages(self._source(root), target)
            self.assertEqual([page.name for page in pages], ["page-001.png", "page-002.png"])
            self.assertEqual([stat.S_IMODE(page.stat().st_mode) for page in pages], [0o644, 0o644])

    def test_pdftoppm_errors_are_rejected_deterministically(self):
        expected_messages = {
            "nonzero": "PDF rasterization failed",
            "timeout": "PDF rasterization failed to start",
            "empty": "converted PDF has no pages",
            "gap": "page sequence is abnormal",
            "stray": "page sequence is abnormal",
        }
        for mode, message in expected_messages.items():
            with self.subTest(mode=mode), tempfile.TemporaryDirectory() as temp:
                root = Path(temp)
                with (
                    patch("docx_page_renderer._load_fitz", side_effect=PageRenderError("fitz missing")),
                    patch("docx_page_renderer._find_pdftoppm", return_value="/fake/pdftoppm"),
                    patch("docx_page_renderer._find_soffice", return_value="/fake/soffice"),
                    patch("docx_page_renderer.subprocess.run", side_effect=self._pdftoppm_run(mode)),
                ):
                    with self.assertRaisesRegex(PageRenderError, message):
                        render_docx_pages(self._source(root), root / "pages")


def _smoke_dependencies_available() -> bool:
    try:
        _find_soffice()
    except PageRenderError:
        return False
    return bool(importlib.util.find_spec("fitz") or shutil.which("pdftoppm"))


@unittest.skipUnless(
    _smoke_dependencies_available(),
    "soffice and either PyMuPDF or pdftoppm are required for the real DOCX smoke test",
)
class DocxPageRendererSmokeTests(unittest.TestCase):
    def test_real_minimal_docx_renders_a_png(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "smoke.docx"
            document = Document()
            document.add_paragraph("DOCX renderer smoke test")
            document.save(source)
            pages = render_docx_pages(source, root / "pages", dpi=72)
            self.assertEqual([path.name for path in pages], ["page-001.png"])
            with Image.open(pages[0]) as image:
                image.verify()


if __name__ == "__main__":
    unittest.main()
