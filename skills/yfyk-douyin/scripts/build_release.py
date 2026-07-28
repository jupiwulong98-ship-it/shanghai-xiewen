#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
WORD_FONT = "Arial Unicode MS"
sys.path.insert(0, str(SCRIPT_DIR))
from card_templates import TEMPLATES, render_cards, wrap_text  # noqa: E402
from inspect_sources import inspect_document  # noqa: E402

def set_run(run, size: float, bold: bool = False, color: str = "222222"):
    run.font.name = WORD_FONT
    rfonts = run._element.get_or_add_rPr().rFonts
    for key in ("eastAsia", "ascii", "hAnsi"):
        rfonts.set(qn(f"w:{key}"), WORD_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str, size=11, bold=False, color="222222", after=6, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if align is not None:
        paragraph.alignment = align
    set_run(paragraph.add_run(text), size, bold, color)
    return paragraph


def ranking_text(item: dict[str, Any]) -> str:
    score = item.get("score")
    base = item["name"]
    if score not in (None, ""):
        base += f"（{score}分）"
    if item.get("description"):
        base += f"：{item['description']}"
    return base


def build_docx(entry: dict[str, Any], cards: list[Path], target: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = WORD_FONT
    normal_rfonts = normal._element.get_or_add_rPr().rFonts
    for key in ("eastAsia", "ascii", "hAnsi"):
        normal_rfonts.set(qn(f"w:{key}"), WORD_FONT)
    normal.font.size = Pt(11)

    add_paragraph(doc, "抖音图文发布稿", 9.5, True, "B45A2A", 4)
    add_paragraph(doc, entry["title"], 22, True, "243447", 12)
    add_paragraph(doc, "正文", 14, True, "B45A2A", 5)
    add_paragraph(doc, entry["intro"], 11, False, "222222", 7)
    if entry["content_mode"] == "ranking":
        add_paragraph(doc, "本次榜单排名如下：", 11, True, "333333", 5)
        items = ((ranking_text(item), "List Number") for item in entry["rankings"])
    else:
        add_paragraph(doc, "核心要点", 11, True, "333333", 5)
        items = ((point, "List Bullet") for point in entry["key_points"])
    for text, style in items:
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.2
        set_run(paragraph.add_run(text), 11)
    add_paragraph(doc, "Tag", 14, True, "B45A2A", 5)
    tags = entry["tags"]
    add_paragraph(doc, " ".join(tags) if isinstance(tags, list) else str(tags), 11, False, "1C6B4F", 12)

    for index, card in enumerate(cards, 1):
        doc.add_page_break()
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.add_run().add_picture(str(card), width=Inches(5.4))
        add_paragraph(doc, f"图 {index}/{len(cards)}", 8.5, False, "777777", 7, WD_ALIGN_PARAGRAPH.CENTER)
    doc.save(target)


def validate_job(job: dict[str, Any]):
    if job.get("version") != 1 or not isinstance(job.get("documents"), list):
        raise ValueError("job.json 必须包含 version=1 和 documents 数组")
    output_filenames: set[str] = set()
    for entry in job["documents"]:
        for field in ("source_path", "output_filename", "title", "intro", "brand", "keyword", "tags", "content_mode", "card_template"):
            if not entry.get(field):
                raise ValueError(f"任务缺少字段 {field}: {entry.get('source_path', '<unknown>')}")
        if entry["card_template"] not in TEMPLATES:
            raise ValueError(f"unknown card_template: {entry['card_template']}")
        tags = entry["tags"]
        if (
            not isinstance(tags, list)
            or not tags
            or any(not isinstance(tag, str) or not tag.strip() for tag in tags)
        ):
            raise ValueError("tags must be a non-empty list of non-empty strings")
        content_mode = entry["content_mode"]
        if content_mode == "ranking":
            rankings = entry.get("rankings")
            if not isinstance(rankings, list) or not rankings:
                raise ValueError("content_mode=ranking requires a non-empty rankings list")
            for item in rankings:
                if not isinstance(item, dict):
                    raise ValueError("each rankings item must be a dict")
                rank = item.get("rank")
                if isinstance(rank, bool) or not isinstance(rank, int) or rank <= 0:
                    raise ValueError("ranking rank must be a positive int")
                if not isinstance(item.get("name"), str) or not item["name"].strip():
                    raise ValueError("ranking name must be a non-empty string")
                if not isinstance(item.get("description"), str):
                    raise ValueError("ranking description must be a string")
                score = item.get("score")
                if (
                    isinstance(score, bool)
                    or score is not None
                    and not isinstance(score, (str, int, float))
                ):
                    raise ValueError("ranking score must be a string, number, or null")
            if "key_points" in entry:
                raise ValueError("content_mode=ranking rejects incompatible key_points field")
        elif content_mode == "article":
            if "rankings" in entry:
                raise ValueError("content_mode=article rejects incompatible rankings field")
            if "ranking_candidate_id" in entry:
                raise ValueError("content_mode=article rejects incompatible ranking_candidate_id field")
            key_points = entry.get("key_points")
            if not isinstance(key_points, list) or not 3 <= len(key_points) <= 5:
                raise ValueError("content_mode=article requires a key_points list of 3 to 5 items")
            if any(not isinstance(point, str) for point in key_points):
                raise ValueError("article key_points must contain only strings")
            if any(not point.strip() for point in key_points):
                raise ValueError("article key_points must contain non-empty strings")
        else:
            raise ValueError(f"unknown content_mode: {content_mode!r}")
        filename = entry["output_filename"]
        if not filename.endswith(".docx") or "/" in filename or "\\" in filename:
            raise ValueError(f"非法输出文件名：{filename}")
        if filename in output_filenames:
            raise ValueError(f"duplicate output_filename: {filename}")
        output_filenames.add(filename)
        if f"#{entry['brand']}" not in tags or f"#{entry['keyword']}" not in tags:
            raise ValueError(f"Tag 必须包含品牌和关键词：{filename}")


def _publish_replace(source: Path, target: Path):
    os.replace(source, target)


def _publish_batch(
    staged_outputs: list[Path],
    targets: list[Path],
    backup_dir: Path,
    overwrite: bool,
):
    backup_dir.mkdir()
    backups: list[tuple[Path, Path]] = []
    published: list[Path] = []
    try:
        for index, target in enumerate(targets):
            if overwrite and target.exists():
                backup = backup_dir / f"{index:06d}.docx"
                _publish_replace(target, backup)
                backups.append((target, backup))
        for staged, target in zip(staged_outputs, targets):
            if not overwrite and target.exists():
                raise FileExistsError(f"目标文件已存在：{target}")
            _publish_replace(staged, target)
            published.append(target)
    except Exception:
        for target in published:
            try:
                target.unlink()
            except FileNotFoundError:
                pass
        for target, backup in reversed(backups):
            if backup.exists():
                os.replace(backup, target)
        raise


def build_release(job_path: Path, output_dir: Path, overwrite: bool = False) -> list[Path]:
    job = json.loads(job_path.read_text(encoding="utf-8"))
    validate_job(job)
    entries = job["documents"]
    for entry in entries:
        source = Path(entry["source_path"])
        if not source.is_file() or source.suffix.lower() != ".docx":
            raise ValueError(f"source_path must be an existing .docx file: {source}")
    targets = [output_dir / entry["output_filename"] for entry in entries]
    invalid_targets = [target for target in targets if target.exists() and not target.is_file()]
    if invalid_targets:
        raise ValueError(
            "existing output target must be a regular file: "
            + ", ".join(map(str, invalid_targets))
        )
    if not overwrite:
        collisions = [target for target in targets if target.exists()]
        if collisions:
            raise FileExistsError(f"目标文件已存在：{', '.join(map(str, collisions))}")

    output_dir.parent.mkdir(parents=True, exist_ok=True)
    staged_outputs: list[Path] = []
    with tempfile.TemporaryDirectory(prefix="docx_to_douyin_", dir=output_dir.parent) as temp:
        temp_root = Path(temp)
        stage_dir = temp_root / "staged"
        stage_dir.mkdir()
        for index, entry in enumerate(entries, 1):
            source = Path(entry["source_path"])
            inspection = inspect_document(source, temp_root / f"inspect-{index:03d}")
            cards = render_cards(
                inspection["blocks"],
                temp_root / f"cards-{index:03d}",
                entry["card_template"],
            )
            staged_target = stage_dir / entry["output_filename"]
            build_docx(entry, cards, staged_target)
            staged_outputs.append(staged_target)

        output_dir.mkdir(parents=True, exist_ok=True)
        _publish_batch(staged_outputs, targets, temp_root / "backups", overwrite)
    return targets


def main() -> int:
    parser = argparse.ArgumentParser(description="Build editable Douyin release DOCX files.")
    parser.add_argument("--job", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--overwrite", action="store_true")
    args = parser.parse_args()
    outputs = build_release(args.job, args.output, args.overwrite)
    print(json.dumps({"outputs": [str(path) for path in outputs]}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
