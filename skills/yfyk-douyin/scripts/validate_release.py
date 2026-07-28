#!/usr/bin/env python3
from __future__ import annotations

import argparse
import io
import json
import re
import tempfile
import zipfile
from xml.etree import ElementTree
from pathlib import Path
from typing import Any

from PIL import Image
from docx import Document
from card_templates import TEMPLATES
from docx_page_renderer import render_docx_pages


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": WORD_NS, "a": DRAWING_NS, "r": REL_NS, "pr": PACKAGE_REL_NS}


def diagnostic(source: str, output: str, code: str, message: str) -> dict[str, str]:
    return {"source_path": source, "output_path": output, "code": code, "message": message}


def media_audit(path: Path) -> tuple[list[tuple[int, int]], list[str], int]:
    sizes: list[tuple[int, int]] = []
    corrupt: list[str] = []
    media_count = 0
    with zipfile.ZipFile(path) as archive:
        for name in sorted(archive.namelist()):
            if not name.startswith("word/media/"):
                continue
            media_count += 1
            try:
                with Image.open(io.BytesIO(archive.read(name))) as image:
                    image.verify()
                    sizes.append(image.size)
            except Exception:
                corrupt.append(name)
    return sizes, corrupt, media_count


def published_card_media(path: Path) -> list[str]:
    """Return media parts used by the builder's page-break, image, caption card structure."""
    with zipfile.ZipFile(path) as archive:
        document = ElementTree.fromstring(archive.read("word/document.xml"))
        relationships = ElementTree.fromstring(archive.read("word/_rels/document.xml.rels"))
    relation_targets = {
        relation.get("Id"): relation.get("Target")
        for relation in relationships.findall("pr:Relationship", NS)
        if relation.get("TargetMode") != "External" and relation.get("Target", "").startswith("media/")
    }
    body = document.find("w:body", NS)
    if body is None:
        return []
    paragraphs = [child for child in body if child.tag == f"{{{WORD_NS}}}p"]
    cards: list[str] = []
    for index, paragraph in enumerate(paragraphs):
        if index == 0 or index + 1 >= len(paragraphs):
            continue
        previous = paragraphs[index - 1]
        following = paragraphs[index + 1]
        has_page_break = any(
            node.get(f"{{{WORD_NS}}}type") == "page"
            for node in previous.findall(".//w:br", NS)
        )
        caption = "".join(node.text or "" for node in following.findall(".//w:t", NS)).strip()
        if not has_page_break or re.fullmatch(r"图 \d+/\d+", caption) is None:
            continue
        for blip in paragraph.findall(".//a:blip", NS):
            target = relation_targets.get(blip.get(f"{{{REL_NS}}}embed"))
            if target:
                cards.append(f"word/{target}")
    return cards


def card_media_audit(path: Path, media_names: list[str]) -> tuple[list[tuple[int, int]], list[str]]:
    sizes: list[tuple[int, int]] = []
    corrupt: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in media_names:
            try:
                with Image.open(io.BytesIO(archive.read(name))) as image:
                    image.verify()
                    sizes.append(image.size)
            except Exception:
                corrupt.append(name)
    return sizes, corrupt


def valid_output_filename(filename: Any) -> bool:
    return (
        isinstance(filename, str)
        and bool(filename)
        and filename not in {".", ".."}
        and "/" not in filename
        and "\\" not in filename
        and not Path(filename).is_absolute()
        and Path(filename).name == filename
        and Path(filename).suffix == ".docx"
    )


def ranking_names(doc: Document) -> list[str]:
    names = []
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name == "List Number":
            text = paragraph.text.strip()
            name = text.split("（", 1)[0].split("：", 1)[0].strip()
            if name:
                names.append(name)
    return names


def styled_texts(doc: Document, style_name: str) -> list[str]:
    return [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.style
        and paragraph.style.name == style_name
        and paragraph.text.strip()
    ]


def validate_release(job_path: Path, output_dir: Path) -> list[dict[str, str]]:
    errors: list[dict[str, str]] = []
    try:
        job = json.loads(job_path.read_text(encoding="utf-8"))
    except Exception as exc:
        return [diagnostic(str(job_path), "", "INVALID_JOB", str(exc))]
    documents = job.get("documents") if isinstance(job, dict) else None
    if not isinstance(documents, list):
        return [diagnostic(str(job_path), "", "INVALID_JOB", "documents 必须是数组")]

    valid_entries: list[dict[str, Any]] = []
    expected_names: set[str] = set()
    seen_names: set[str] = set()
    for index, entry in enumerate(documents):
        if not isinstance(entry, dict):
            errors.append(diagnostic(str(job_path), "", "MALFORMED_ENTRY", f"documents[{index}] 必须是对象"))
            continue
        template_id = entry.get("card_template")
        if not template_id:
            errors.append(diagnostic(
                str(entry.get("source_path", job_path)),
                "",
                "MISSING_CARD_TEMPLATE",
                f"documents[{index}] 缺少 card_template",
            ))
        elif template_id not in TEMPLATES:
            errors.append(diagnostic(
                str(entry.get("source_path", job_path)),
                "",
                "UNKNOWN_CARD_TEMPLATE",
                str(template_id),
            ))
        filename = entry.get("output_filename")
        if not valid_output_filename(filename):
            errors.append(diagnostic(
                str(entry.get("source_path", job_path)),
                "",
                "INVALID_OUTPUT_FILENAME",
                f"documents[{index}] 的 output_filename 非法：{filename!r}",
            ))
            continue
        if filename in seen_names:
            errors.append(diagnostic(
                str(entry.get("source_path", job_path)),
                str(output_dir / filename),
                "DUPLICATE_OUTPUT_FILENAME",
                f"重复 output_filename：{filename}",
            ))
            continue
        seen_names.add(filename)
        expected_names.add(filename)
        valid_entries.append(entry)

    if output_dir.exists() and not output_dir.is_dir():
        errors.append(diagnostic(
            str(job_path),
            str(output_dir),
            "OUTPUT_DIR_NOT_DIRECTORY",
            "输出路径存在但不是目录",
        ))
        return errors
    try:
        actual_files = {path.name for path in output_dir.iterdir()} if output_dir.exists() else set()
    except OSError as exc:
        errors.append(diagnostic(
            str(job_path),
            str(output_dir),
            "OUTPUT_DIR_UNREADABLE",
            str(exc),
        ))
        return errors
    for extra in sorted(actual_files - expected_names):
        errors.append(diagnostic("", str(output_dir / extra), "EXTRA_OUTPUT", "输出目录只能包含最终 DOCX"))
    for missing in sorted(expected_names - actual_files):
        errors.append(diagnostic("", str(output_dir / missing), "MISSING_OUTPUT", "缺少预期 DOCX"))

    for entry in valid_entries:
        source = str(entry.get("source_path", ""))
        output = output_dir / entry["output_filename"]
        if not output.exists():
            continue
        try:
            doc = Document(output)
        except Exception as exc:
            errors.append(diagnostic(source, str(output), "INVALID_DOCX", str(exc)))
            continue
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        title = entry.get("title")
        required_texts = ((title, "MISSING_TITLE"), ("正文", "MISSING_BODY"), ("Tag", "MISSING_TAG"))
        for required, code in required_texts:
            if not isinstance(required, str) or not required or required not in texts:
                errors.append(diagnostic(source, str(output), code, f"缺少：{required}"))
        if "榜单" in texts:
            errors.append(diagnostic(source, str(output), "PEER_RANKING_HEADING", "榜单不能与正文平级"))
        hashtags = set(re.findall(r"(?<!\S)#[^\s]+", " ".join(texts)))
        brand = entry.get("brand")
        keyword = entry.get("keyword")
        if not isinstance(brand, str) or f"#{brand}" not in hashtags:
            errors.append(diagnostic(source, str(output), "MISSING_BRAND_TAG", str(brand)))
        if not isinstance(keyword, str) or f"#{keyword}" not in hashtags:
            errors.append(diagnostic(source, str(output), "MISSING_KEYWORD_TAG", str(keyword)))

        mode = entry.get("content_mode")
        if mode == "ranking":
            if "本次榜单排名如下：" not in texts:
                errors.append(diagnostic(
                    source,
                    str(output),
                    "MISSING_RANKING_LEAD",
                    "缺少：本次榜单排名如下：",
                ))
            rankings = entry.get("rankings")
            if not isinstance(rankings, list) or any(not isinstance(item, dict) for item in rankings):
                errors.append(diagnostic(source, str(output), "MALFORMED_ENTRY", "rankings 必须是对象数组"))
            else:
                expected_ranking = [
                    item.get("name")
                    for item in rankings
                    if isinstance(item.get("name"), str)
                ]
                actual_ranking = ranking_names(doc)
                if len(expected_ranking) != len(rankings):
                    errors.append(diagnostic(source, str(output), "MALFORMED_ENTRY", "ranking 缺少有效 name"))
                elif actual_ranking != expected_ranking:
                    errors.append(diagnostic(
                        source,
                        str(output),
                        "RANKING_MISMATCH",
                        f"expected={expected_ranking}; actual={actual_ranking}",
                    ))
            bullet_texts = styled_texts(doc, "List Bullet")
            if bullet_texts:
                errors.append(diagnostic(
                    source,
                    str(output),
                    "RANKING_HAS_KEY_POINTS",
                    f"ranking 模式不能包含项目符号要点：{bullet_texts}",
                ))
        elif mode == "article":
            expected_points = entry.get("key_points")
            actual_points = styled_texts(doc, "List Bullet")
            if not isinstance(expected_points, list) or any(not isinstance(point, str) for point in expected_points):
                errors.append(diagnostic(source, str(output), "MALFORMED_ENTRY", "key_points 必须是字符串数组"))
            elif actual_points != expected_points:
                errors.append(diagnostic(
                    source,
                    str(output),
                    "KEY_POINTS_MISMATCH",
                    f"expected={expected_points}; actual={actual_points}",
                ))
            if "本次榜单排名如下：" in texts:
                errors.append(diagnostic(
                    source,
                    str(output),
                    "ARTICLE_HAS_RANKING_COPY",
                    "article 模式不能包含榜单引导语",
                ))
            numbered = styled_texts(doc, "List Number")
            if numbered:
                errors.append(diagnostic(
                    source,
                    str(output),
                    "ARTICLE_NUMBERED_LIST",
                    f"article 模式不能包含编号列表：{numbered}",
                ))
        else:
            errors.append(diagnostic(
                source,
                str(output),
                "INVALID_CONTENT_MODE",
                f"未知或缺失 content_mode：{mode!r}",
            ))

        try:
            sizes, corrupt_media, media_count = media_audit(output)
            card_media = published_card_media(output)
            card_sizes, corrupt_cards = card_media_audit(output, card_media)
        except (OSError, KeyError, zipfile.BadZipFile, ElementTree.ParseError) as exc:
            errors.append(diagnostic(source, str(output), "INVALID_DOCX", str(exc)))
            continue
        if not card_media:
            errors.append(diagnostic(source, str(output), "NO_CARDS", "DOCX 未嵌入图片"))
        for name in corrupt_media:
            errors.append(diagnostic(
                source,
                str(output),
                "CORRUPT_MEDIA",
                f"无法读取嵌入媒体：{name}",
            ))
        for name in corrupt_cards:
            if name not in corrupt_media:
                errors.append(diagnostic(
                    source,
                    str(output),
                    "CORRUPT_MEDIA",
                    f"无法读取嵌入媒体：{name}",
                ))
        for size in card_sizes:
            if size != (1080, 1440):
                errors.append(diagnostic(source, str(output), "WRONG_CARD_SIZE", str(size)))
                errors.append(diagnostic(source, str(output), "INVALID_CARD_SIZE", str(size)))

        try:
            with tempfile.TemporaryDirectory(prefix="validate_source_pages_") as temp:
                expected_cards = render_docx_pages(Path(source), Path(temp) / "pages")
        except Exception as exc:
            errors.append(diagnostic(
                source,
                str(output),
                "SOURCE_PAGE_RENDER_FAILED",
                str(exc),
            ))
        else:
            if len(card_media) != len(expected_cards):
                errors.append(diagnostic(
                    source,
                    str(output),
                    "CARD_COUNT_MISMATCH",
                    f"expected={len(expected_cards)}; actual={len(card_media)}",
                ))
    return errors


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate Douyin release DOCX files.")
    parser.add_argument("--job", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    errors = validate_release(args.job, args.output)
    print(json.dumps({"valid": not errors, "errors": errors}, ensure_ascii=False, indent=2))
    return 0 if not errors else 1


if __name__ == "__main__":
    raise SystemExit(main())
